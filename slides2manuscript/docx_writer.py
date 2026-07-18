"""생성된 원고를 드라이한(꾸미지 않은) docx로 저장한다.

서식 원칙: 본문 글꼴 하나, 장 제목만 약간 키움, 색·표·테두리·머리글 없음.
한국어 A4 기준으로 분량 계산이 맞도록 글꼴/줄간격/여백을 고정한다.

공식 규격(전문가 활용비 강의료 및 원고료 책정표):
- A4 1페이지: 12Font / 35Line, 신명조
- 상하 여백 15mm, 좌우 여백 20mm, 머리말·꼬리말 15mm

원고 구성: 제목 페이지 → 목차 → 본문(장별 페이지 분리) → 참고자료.
목차의 페이지 번호는 2-pass 렌더링(1차 저장 → PDF 변환으로 실측 → 페이지 번호
채워 2차 재저장)으로 채운다.
"""

from __future__ import annotations

import os
import platform

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.shared import Cm, Mm, Pt
from docx.oxml.ns import qn

from . import font_setup
from .generate import Section


def _default_body_font() -> str:
    """생성 시점 OS에 실제로 존재하는 명조체 폰트 이름을 반환한다.

    규격은 "신명조"지만 이 이름의 폰트가 시스템에 없으면 LibreOffice·Word가
    엉뚱한 산세리프로 대체해 버려 명조체 규격이 깨진다. 그래서:
    - 먼저 나눔명조(OFL, 자동 설치 가능)가 있으면 그것을 쓰고,
    - 없으면 각 OS 기본 명조체로 폴백한다.
    `--font` 로 override 가능.
    """
    if font_setup.is_font_installed("NanumMyeongjo"):
        return "나눔명조"
    system = platform.system()
    if system == "Darwin":
        if os.path.exists("/System/Library/Fonts/Supplemental/AppleMyungjo.ttf"):
            return "AppleMyungjo"
    if system == "Windows":
        return "바탕"
    return "NanumMyeongjo"


BODY_FONT = _default_body_font()
BODY_SIZE = 12
HEADING_SIZE = 13
TITLE_SIZE = 15
COVER_TITLE_SIZE = 22       # 표지 페이지의 강의 제목 크기
COVER_AUTHOR_SIZE = 14      # 표지의 강사명 크기
COVER_DATE_SIZE = 12        # 표지의 날짜 크기
# A4(297mm) - 상하 여백 30mm = 267mm 안에 35줄이 들어가도록 줄 높이를 EXACTLY로 고정한다.
LINE_HEIGHT_PT = 21.0


def _apply_rfonts(rpr, name: str) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _set_korean_font(run, name: str) -> None:
    run.font.name = name
    _apply_rfonts(run._element.get_or_add_rPr(), name)


def _base_style(doc: Document, body_font: str) -> None:
    style = doc.styles["Normal"]
    style.font.name = body_font
    style.font.size = Pt(BODY_SIZE)
    _apply_rfonts(style.element.get_or_add_rPr(), body_font)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(LINE_HEIGHT_PT)
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.header_distance = Mm(15)
        section.footer_distance = Mm(15)


def _add_paragraph(doc: Document, text: str, size: int, bold: bool, body_font: str,
                   space_before: int = 0, align=None):
    p = doc.add_paragraph()
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    _set_korean_font(run, body_font)
    return p


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _write_title_page(doc: Document, doc_title: str, author: str | None,
                       date: str | None, body_font: str) -> None:
    """표지 페이지: 강의 제목, 강사, 날짜. 중앙 정렬로 여백 있게 배치."""
    # 페이지 상단에 여백 (빈 줄 여러 개)
    for _ in range(8):
        doc.add_paragraph()
    # 강의 제목
    _add_paragraph(doc, doc_title, COVER_TITLE_SIZE, bold=True, body_font=body_font,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    # 제목 아래 여백
    for _ in range(5):
        doc.add_paragraph()
    # 강사명
    if author:
        _add_paragraph(doc, author, COVER_AUTHOR_SIZE, bold=False, body_font=body_font,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()
    # 날짜
    if date:
        _add_paragraph(doc, date, COVER_DATE_SIZE, bold=False, body_font=body_font,
                       align=WD_ALIGN_PARAGRAPH.CENTER)


def _write_toc(doc: Document, sections: list[Section], page_map: dict[int, int] | None,
                body_font: str) -> None:
    """목차 페이지. page_map이 있으면 각 장 옆에 페이지 번호를 표시."""
    _add_paragraph(doc, "목차", TITLE_SIZE, bold=True, body_font=body_font,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    for i, sec in enumerate(sections):
        num = f"{i + 1}. "
        # 페이지 번호가 있으면 오른쪽 정렬로 붙임(간단히 공백 padding). 없으면 제목만.
        if page_map and i in page_map:
            # 제목과 페이지 사이를 " ......... " 으로 채워 시각적 정렬
            title_text = num + sec.title
            page_text = str(page_map[i])
            # 대략 60자 폭에 맞춰 dot leader
            dots_len = max(3, 60 - len(title_text) - len(page_text))
            line = f"{title_text} {'·' * dots_len} {page_text}"
        else:
            line = num + sec.title
        _add_paragraph(doc, line, BODY_SIZE, bold=False, body_font=body_font)


def _write_body_section(doc: Document, sec: Section, body_font: str) -> None:
    """장 제목 + 본문. 장 제목은 볼드로 강조."""
    _add_paragraph(doc, sec.title, HEADING_SIZE, bold=True, body_font=body_font)
    doc.add_paragraph()
    for para in _split_paragraphs(sec.body):
        _add_paragraph(doc, para, BODY_SIZE, bold=False, body_font=body_font)


def _write_references(doc: Document, references: list[str], body_font: str) -> None:
    """참고자료 페이지. 각 항목을 문단으로 나열."""
    _add_paragraph(doc, "참고자료", TITLE_SIZE, bold=True, body_font=body_font,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    for i, ref in enumerate(references, 1):
        _add_paragraph(doc, f"{i}. {ref}", BODY_SIZE, bold=False, body_font=body_font)


def _find_section_pages(pdf_path: str, sections: list[Section],
                         min_page: int = 1) -> dict[int, int]:
    """PDF 안에서 각 장 제목의 시작 페이지를 찾아 매핑을 반환한다.

    min_page 이후의 페이지만 검색해 목차 페이지 안의 제목 등장을 배제한다.
    반환: {section_index: 1-based page number}
    """
    try:
        import fitz
    except ImportError:
        return {}
    page_map: dict[int, int] = {}
    doc = fitz.open(pdf_path)
    try:
        # 각 페이지 텍스트를 미리 뽑아 검색 속도 확보
        page_texts = [(i + 1, p.get_text("text")) for i, p in enumerate(doc)]
    finally:
        doc.close()
    used = set()
    for i, sec in enumerate(sections):
        for page_num, text in page_texts:
            if page_num < min_page or page_num in used:
                continue
            if sec.title in text:
                page_map[i] = page_num
                used.add(page_num)
                break
    return page_map


def write_docx(
    out_path: str,
    doc_title: str,
    sections: list[Section],
    body_font: str | None = None,
    author: str | None = None,
    date: str | None = None,
    title_page: bool = True,
    toc: bool = True,
    references: list[str] | None = None,
) -> None:
    """docx 파일을 생성한다.

    - title_page: 표지 페이지 (강의 제목·강사·날짜)
    - toc: 목차 페이지. True 면 두 번 렌더링해 각 장의 실제 시작 페이지 번호를 채운다.
    - references: 참고자료 문단 리스트. None 이면 참고자료 페이지 생략.
    """
    body_font = body_font or BODY_FONT

    def _build(page_map: dict[int, int] | None) -> None:
        doc = Document()
        _base_style(doc, body_font)

        if title_page:
            _write_title_page(doc, doc_title, author, date, body_font)
            _add_page_break(doc)
        if toc:
            _write_toc(doc, sections, page_map, body_font)
            _add_page_break(doc)
        for i, sec in enumerate(sections):
            if i > 0:
                _add_page_break(doc)
            _write_body_section(doc, sec, body_font)
        if references:
            _add_page_break(doc)
            _write_references(doc, references, body_font)

        doc.save(out_path)

    # 1차: 목차의 페이지 번호를 알기 전 상태로 저장
    _build(page_map=None)

    # 목차가 있으면 실제 페이지 매핑을 얻어 2차 재저장
    if toc:
        from . import pdf_export  # 순환 import 회피
        pdf_path, _info = pdf_export.convert_to_pdf(out_path)
        if pdf_path and os.path.isfile(pdf_path):
            # 표지 다음 페이지가 목차 시작. 목차 자체는 1~2쪽 정도라 min_page를 여유롭게 잡는다.
            min_page = (2 if title_page else 1)  # 목차 첫 페이지
            page_map = _find_section_pages(pdf_path, sections, min_page=min_page + 1)
            if page_map:
                _build(page_map=page_map)


def _split_paragraphs(body: str) -> list[str]:
    """본문 텍스트를 문단 단위로 나눈다."""
    chunks = [c.strip() for c in body.replace("\r\n", "\n").split("\n\n")]
    out: list[str] = []
    for c in chunks:
        if not c:
            continue
        out.append(" ".join(line.strip() for line in c.split("\n") if line.strip()))
    return out


def manuscript_char_count(sections: list[Section]) -> int:
    return sum(len(s.body) for s in sections)
